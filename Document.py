from docx import Document

from DayTask import DayTask


class DOC:
    def __init__(self, start_week, endweek):
        self.document = Document()
        self.document.add_heading(f'Plan nauki na tydzień od {start_week} do {endweek}', 0)
    def AddTabet(self,  records: list[DayTask] | DayTask, data):
        table = self.document.add_table(rows=1, cols=2)
        if type(records) == list:
            for daytask in records:
                row = table.add_row().cells
                row[0].text = str(data)
                row[1].text  = daytask.getInformation()
        else:
            row = table.add_row().cells
            row[0].text = data
            row[1].text = records.getInformation()

    def Save(self):
        self.document.save('plan.docx')
